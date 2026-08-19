#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown → DOCX converter for report formatting rules.

Usage:
  python md_to_docx.py INPUT.md OUTPUT.docx

Default format rules (docx 格式要求.txt):
  1. 中文宋体 / 英文 Times New Roman
  2. 正文小四 (12pt)
  3. 表格、代码五号 (10.5pt)
  4. 页边距窄 (1.27cm)
  5. 表格适应窗口
  6. 行距 1.5
  7. 页码底部居中，五号
  8. 插图用地址代替（仍写【插图地址】行）
  9. 自动生成目录（TOC 域）
  10. 标题使用 Heading 样式以便进入目录
  11. 表格和图片前后各 0.5 行间隔
  12. 图及图名称居中
  13. [文字](url) → Word 超链接
  14. 超链接文字 RGB(0,102,204)，宋体/Times New Roman，字号与上下文一致
  15. 插图地址编号为 assets/N.filename；并尝试嵌入缩放图片
  16. 「…」→“… ”
  17. LaTeX 公式转为 Word 公式（OMML）
  18. 代码块整段黑底白字
  19. 适合流程图的内容：先 Mermaid→PNG（见 mermaid_to_png.py），再按 §15 嵌入

版式样例：STAGE3.docx（Heading 1–4、居中图注、真超链接）。
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

SIZE_BODY, SIZE_TABLE, SIZE_CODE = Pt(12), Pt(10.5), Pt(10.5)
SIZE_H1, SIZE_H2, SIZE_H3, SIZE_H4 = Pt(18), Pt(16), Pt(14), Pt(12)
SIZE_PAGE = Pt(10.5)
FONT_CN, FONT_EN, FONT_CODE = "宋体", "Times New Roman", "Consolas"
NARROW_CM = 1.27
# 0.5 行 = 小四 12pt × 1.5 行距 × 0.5
HALF_LINE = Pt(9)
# docx 格式要求.txt §14：RGB(0,102,204)
LINK_BLUE = "0066CC"
# 嵌入图默认最大宽（窄页边距下可用宽约 18cm，取 14cm 留白）
IMG_MAX_WIDTH = Cm(14)
IMG_MAX_HEIGHT = Cm(18)
CODE_BG = "000000"
CODE_FG = RGBColor(255, 255, 255)
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
CAPTION_RE = re.compile(r"^\*?(图[：:].+|示意[：:].+)\*?$")
MD_IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)")
# §17: delimited math first, then LaTeX-like fragments
DELIM_MATH_RE = re.compile(
    r"\$\$(.+?)\$\$|\\\[(.+?)\\\]|\\\((.+?)\\\)|(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)",
    re.S,
)
BARE_MATH_RE = re.compile(
    r"\([^()]*(?:\^|_\{|\\[a-zA-Z]+)[^()]*\)"
    r"|[A-Za-z][A-Za-z0-9]*(?:\^\{[^}]*\}|\^[^\s{,_]|_\{[^}]*\})+"
    r"|\\(?:neq|ne|leq|geq|times|cdot|pm|infty|rightarrow|leftarrow|in|subset|approx|equiv)(?![A-Za-z])"
)
LATEX_SYMBOLS = {
    "neq": "≠",
    "ne": "≠",
    "leq": "≤",
    "geq": "≥",
    "times": "×",
    "cdot": "·",
    "pm": "±",
    "infty": "∞",
    "rightarrow": "→",
    "leftarrow": "←",
    "Rightarrow": "⇒",
    "in": "∈",
    "subset": "⊂",
    "approx": "≈",
    "equiv": "≡",
    "ast": "*",
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "theta": "θ",
    "lambda": "λ",
    "mu": "μ",
    "pi": "π",
    "sigma": "σ",
    "phi": "φ",
    "omega": "ω",
    "sum": "∑",
    "prod": "∏",
    "ldots": "…",
    "cdots": "⋯",
}


def normalize_quotes(text: str) -> str:
    """docx 格式要求 §16：强调引号「…」→中文弯引号 “… ”。"""
    return text.replace("「", "“").replace("」", "”")


def _me(tag: str):
    return etree.Element("{%s}%s" % (MATH_NS, tag))


def _m_val(el, val: str):
    el.set("{%s}val" % MATH_NS, val)


def _math_run(text: str, italic: bool | None = None):
    """Build m:r. italic=None → auto (letters italic, others plain)."""
    r = _me("r")
    rPr = _me("rPr")
    sty = _me("sty")
    if italic is None:
        italic = bool(text) and text.isalpha()
    _m_val(sty, "i" if italic else "p")
    rPr.append(sty)
    r.append(rPr)
    t = _me("t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _math_join(nodes: list):
    if len(nodes) == 1:
        return nodes[0]
    e = _me("e")
    for n in nodes:
        e.append(n)
    return e


class _LatexParser:
    def __init__(self, s: str):
        self.s = s.strip()
        self.i = 0

    def peek(self) -> str:
        return self.s[self.i] if self.i < len(self.s) else ""

    def parse_expr(self) -> list:
        nodes = []
        while self.i < len(self.s) and self.peek() not in ")}":
            before = self.i
            nodes.append(self.parse_scripted())
            if self.i <= before:
                break
        return nodes or [_math_run("")]

    def parse_group(self) -> list:
        if self.peek() == "{":
            self.i += 1
            nodes = self.parse_expr()
            if self.peek() == "}":
                self.i += 1
            return nodes
        # Unbraced ^/_ takes one atom so a^*_{t,1} becomes sSubSup, not * with a subscript.
        return [self.parse_atom()]

    def parse_atom(self):
        ch = self.peek()
        if not ch:
            return _math_run("")
        if ch == "{":
            return _math_join(self.parse_group())
        if ch == "\\":
            return self.parse_command()
        if ch in "()[]":
            self.i += 1
            return _math_run(ch, italic=False)
        if ch.isdigit() or ch in "=+-/*<>,.|:'":
            j = self.i
            while self.i < len(self.s) and (self.s[self.i].isdigit() or self.s[self.i] in ".,"):
                self.i += 1
            if self.i == j:
                self.i += 1
                return _math_run(ch, italic=False)
            return _math_run(self.s[j : self.i], italic=False)
        if ch.isalpha():
            j = self.i
            while self.i < len(self.s) and self.s[self.i].isalpha():
                self.i += 1
            return _math_run(self.s[j : self.i], italic=True)
        self.i += 1
        return _math_run(ch, italic=False)

    def parse_command(self):
        self.i += 1
        if self.i < len(self.s) and not self.s[self.i].isalpha():
            cmd = self.s[self.i]
            self.i += 1
            if cmd in "{}":
                return _math_run(cmd, italic=False)
            return _math_run(cmd, italic=False)
        j = self.i
        while self.i < len(self.s) and self.s[self.i].isalpha():
            self.i += 1
        cmd = self.s[j : self.i]
        if cmd in ("left", "right"):
            return self.parse_atom()
        if cmd == "frac":
            f = _me("f")
            num = _me("num")
            den = _me("den")
            for n in self.parse_group():
                num.append(n)
            for n in self.parse_group():
                den.append(n)
            f.extend([num, den])
            return f
        if cmd in ("sqrt",):
            rad = _me("rad")
            deg = _me("deg")
            e = _me("e")
            for n in self.parse_group():
                e.append(n)
            rad.extend([deg, e])
            return rad
        if cmd in ("mathrm", "text", "textbf"):
            return _math_join(self.parse_group())
        if cmd in LATEX_SYMBOLS:
            return _math_run(LATEX_SYMBOLS[cmd], italic=False)
        return _math_run("\\" + cmd, italic=False)

    def parse_scripted(self):
        base = self.parse_atom()
        sub = sup = None
        # Use a tuple: "" in "^_" is True in Python and would loop forever at EOS.
        while self.peek() in ("^", "_"):
            mark = self.peek()
            self.i += 1
            grp = self.parse_group()
            if mark == "_":
                sub = grp
            else:
                sup = grp
        if sub is None and sup is None:
            return base
        if sub is not None and sup is not None:
            node = _me("sSubSup")
            e, sub_el, sup_el = _me("e"), _me("sub"), _me("sup")
            e.append(base)
            for n in sub:
                sub_el.append(n)
            for n in sup:
                sup_el.append(n)
            node.extend([e, sub_el, sup_el])
            return node
        if sub is not None:
            node = _me("sSub")
            e, sub_el = _me("e"), _me("sub")
            e.append(base)
            for n in sub:
                sub_el.append(n)
            node.extend([e, sub_el])
            return node
        node = _me("sSup")
        e, sup_el = _me("e"), _me("sup")
        e.append(base)
        for n in sup:
            sup_el.append(n)
        node.extend([e, sup_el])
        return node


def latex_to_omml(latex: str, display: bool = False):
    """§17: convert LaTeX to Word OMML (m:oMath / m:oMathPara)."""
    raw = latex.strip()
    raw = raw.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    om = etree.Element("{%s}oMath" % MATH_NS, nsmap={"m": MATH_NS})
    try:
        nodes = _LatexParser(raw).parse_expr()
        for n in nodes:
            om.append(n)
    except Exception:
        om.append(_math_run(raw, italic=False))
    if display:
        para = etree.Element("{%s}oMathPara" % MATH_NS, nsmap={"m": MATH_NS})
        para.append(om)
        return para
    return om


def _merge_spans(spans: list[tuple[int, int, str, bool]], text: str):
    if not spans:
        return []
    spans = sorted(spans)
    out = [list(spans[0])]
    for s, e, latex, disp in spans[1:]:
        prev = out[-1]
        gap = text[prev[1] : s]
        raw_prev = text[prev[0] : prev[1]] == prev[2]
        raw_cur = text[s:e] == latex
        if (not disp) and (not prev[3]) and raw_prev and raw_cur and re.fullmatch(r"[\s\\]*", gap or ""):
            prev[1] = e
            prev[2] = text[prev[0] : e]
        else:
            out.append([s, e, latex, disp])
    return [(a, b, c, d) for a, b, c, d in out]


def extract_math_segments(text: str) -> list[tuple[str, str, bool]]:
    """Yield (kind, payload, display) where kind is 'text' or 'math'."""
    spans = []
    for m in DELIM_MATH_RE.finditer(text):
        latex = next(g for g in m.groups() if g is not None)
        display = m.group(1) is not None or m.group(2) is not None
        spans.append((m.start(), m.end(), latex, display))
    occupied = [False] * (len(text) + 1)
    for s, e, *_ in spans:
        for k in range(s, e):
            occupied[k] = True
    for m in BARE_MATH_RE.finditer(text):
        if any(occupied[k] for k in range(m.start(), m.end())):
            continue
        spans.append((m.start(), m.end(), m.group(0), False))
    spans = _merge_spans(spans, text)
    segs = []
    pos = 0
    for s, e, latex, disp in spans:
        if s > pos:
            segs.append(("text", text[pos:s], False))
        segs.append(("math", latex, disp))
        pos = e
    if pos < len(text):
        segs.append(("text", text[pos:], False))
    return segs or [("text", text, False)]


def append_omml(paragraph, latex: str, display: bool = False):
    paragraph._p.append(latex_to_omml(latex, display=display))
    if display:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


class ImageIndex:
    """Assign sequential assets/N.filename display paths; resolve disk files."""

    def __init__(self, md_dir: Path | None = None):
        self.n = 0
        self.md_dir = Path(md_dir) if md_dir else Path.cwd()

    def numbered(self, src: str) -> str:
        name = Path(str(src).replace("\\", "/")).name
        name = re.sub(r"^\d+\.", "", name)
        self.n += 1
        return f"assets/{self.n}.{name}"

    def resolve(self, src: str) -> Path | None:
        """Locate image on disk relative to the Markdown file."""
        raw = str(src).strip().replace("\\", "/")
        candidates: list[Path] = []
        p = Path(raw)
        if p.is_absolute():
            candidates.append(p)
        else:
            candidates.append(self.md_dir / raw)
            candidates.append(self.md_dir / "assets" / p.name)
            # e.g. assets/sam2grasp/fig1.png already under md_dir
            if raw.startswith("./"):
                candidates.append(self.md_dir / raw[2:])
        for c in candidates:
            if c.is_file():
                return c.resolve()
        return None


def set_run_font(run, size_pt, bold=False, italic=False, code=False, color=None):
    run.bold, run.italic, run.font.size = bold, italic, size_pt
    name = FONT_CODE if code else FONT_EN
    run.font.name = name
    rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, space_after=Pt(6), first_line=False, left_indent=None, space_before=Pt(0)):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after, pf.space_before = space_after, space_before
    if first_line:
        pf.first_line_indent = Cm(0.74)
    if left_indent is not None:
        pf.left_indent = left_indent


def configure_heading_style(style, size_pt, space_before=Pt(12), space_after=Pt(6)):
    style.font.name, style.font.size, style.font.bold = FONT_EN, size_pt, True
    style.font.color.rgb = RGBColor(0, 0, 0)
    rFonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before, pf.space_after = space_before, space_after


def set_cell_shading(cell, fill="D9E2F3"):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def set_table_autofit_window(table):
    tbl, tblPr = table._tbl, table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for tag in ("w:tblW", "w:tblLayout"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)
    table.autofit = True
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = True


def _pt_half(size_pt) -> str:
    return str(int(float(size_pt.pt) * 2))


def add_hyperlink(paragraph, text, url, size=SIZE_BODY, bold=False):
    """Word hyperlink: visible text is `text`, target is `url`.

    Color RGB(0,102,204); 宋体 / Times New Roman; size matches surrounding context
    (docx 格式要求.txt §14). No forced underline (aligned with STAGE3.docx).
    """
    url = url.strip()
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_EN)
    rFonts.set(qn("w:hAnsi"), FONT_EN)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), _pt_half(size))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), _pt_half(size))
    rPr.append(szCs)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    rPr.append(color)
    if bold:
        rPr.append(OxmlElement("w:b"))

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_formatted_runs(paragraph, text, size=SIZE_BODY, base_bold=False):
    text = normalize_quotes(text)
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size, bold=base_bold)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(normalize_quotes(token[2:-2]))
            set_run_font(run, size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, SIZE_CODE if size == SIZE_BODY else size, code=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size, bold=base_bold)


def add_text_with_math(paragraph, text, size=SIZE_BODY, base_bold=False):
    """Render text, converting LaTeX fragments to Word OMML (§17)."""
    for kind, payload, display in extract_math_segments(text):
        if kind == "math":
            append_omml(paragraph, payload, display=display)
        elif payload:
            add_formatted_runs(paragraph, payload, size, base_bold)


def add_inline_runs(paragraph, text, size=SIZE_BODY, base_bold=False):
    """Render markdown inline: **bold**, `code`, LaTeX math, [text](url)."""
    text = normalize_quotes(text)
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            add_text_with_math(paragraph, text[pos : m.start()], size, base_bold)
        display, url = m.group(1), m.group(2)
        bold = base_bold or (display.startswith("**") and display.endswith("**"))
        display_clean = normalize_quotes(display.replace("**", "").replace("`", ""))
        add_hyperlink(paragraph, display_clean, url, size=size, bold=bold)
        pos = m.end()
    if pos < len(text):
        add_text_with_math(paragraph, text[pos:], size, base_bold)


def add_heading_styled(doc, text, level):
    text = text.strip()
    style_map = {0: 1, 1: 1, 2: 2, 3: 3, 4: 4, 5: 4}
    hlevel = style_map.get(level, 4)
    p = doc.add_paragraph(style=f"Heading {hlevel}")
    if p.runs:
        p.runs[0].text = ""
    sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3, 4: SIZE_H4}
    add_inline_runs(p, text, size=sizes[hlevel], base_bold=True)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_body_paragraph(doc, text, first_line=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    add_inline_runs(p, text, size=SIZE_BODY)
    return p


def add_list_item(doc, text, ordered=False, index=1):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False, left_indent=Cm(0.74))
    run = p.add_run(f"{index}. " if ordered else "• ")
    set_run_font(run, SIZE_BODY)
    add_inline_runs(p, text, size=SIZE_BODY)
    return p


def add_code_block(doc, code_text):
    """§18: fenced code → black paragraph shading, white Consolas."""
    lines = code_text.splitlines() or [""]
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = HALF_LINE if i == 0 else Pt(0)
        pf.space_after = HALF_LINE if i == len(lines) - 1 else Pt(0)
        pf.left_indent = Cm(0.5)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CODE_BG)
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)
        run = p.add_run(line if line else " ")
        set_run_font(run, SIZE_CODE, code=True, color=CODE_FG)
    sp = doc.add_paragraph()
    set_paragraph_format(sp, space_after=Pt(6))


def add_half_line_gap(doc):
    """表格、图片前后的 0.5 行空白（小四 × 1.5 行距的一半）。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = HALF_LINE
    pf.space_before = pf.space_after = Pt(0)
    run = p.add_run(" ")
    set_run_font(run, HALF_LINE)


def add_table(doc, rows):
    if not rows:
        return
    add_half_line_gap(doc)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_autofit_window(table)
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(2)
            raw = row[j] if j < len(row) else ""
            add_inline_runs(p, raw, size=SIZE_TABLE, base_bold=(i == 0))
            if i == 0:
                set_cell_shading(cell)
    add_half_line_gap(doc)


def _picture_size_cm(path: Path) -> tuple[Cm, Cm | None]:
    """Pick width (and optional height) so image fits page without huge height."""
    try:
        from PIL import Image  # optional; fallback if missing

        with Image.open(path) as im:
            w_px, h_px = im.size
        if w_px <= 0 or h_px <= 0:
            return IMG_MAX_WIDTH, None
        aspect = h_px / w_px
        width = IMG_MAX_WIDTH
        height = Cm(width.cm * aspect)
        if height > IMG_MAX_HEIGHT:
            height = IMG_MAX_HEIGHT
            width = Cm(height.cm / aspect)
        return width, height
    except Exception:
        return IMG_MAX_WIDTH, None


def add_image(doc, display_src: str, file_path: Path | None, caption: str | None = None):
    """§8+§15: centered address line, then try embed scaled picture; caption centered."""
    add_half_line_gap(doc)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【插图地址】{display_src}")
    set_run_font(run, SIZE_BODY)

    if file_path is not None and file_path.is_file():
        pic_p = doc.add_paragraph()
        set_paragraph_format(pic_p, first_line=False, space_after=Pt(0))
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic_p.add_run()
        width, height = _picture_size_cm(file_path)
        try:
            if height is not None:
                pic_run.add_picture(str(file_path), width=width, height=height)
            else:
                pic_run.add_picture(str(file_path), width=width)
        except Exception as exc:
            warn = pic_p.add_run(f"（嵌入失败: {exc}）")
            set_run_font(warn, SIZE_TABLE, italic=True)
    elif file_path is None:
        print(f"WARNING: image not found for {display_src}", file=sys.stderr)

    if caption:
        cap = normalize_quotes(re.sub(r"</?em>", "", caption).strip().strip("*").strip())
        cp = doc.add_paragraph()
        set_paragraph_format(cp, first_line=False)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(cap)
        set_run_font(run, SIZE_TABLE, italic=True)
    add_half_line_gap(doc)


# 兼容旧名
def add_image_as_address(doc, src: str, caption: str | None = None, file_path: Path | None = None):
    add_image(doc, src, file_path, caption=caption)


def add_caption(doc, text: str):
    cap = normalize_quotes(text.strip().strip("*").strip())
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cap)
    set_run_font(run, SIZE_TABLE, italic=True)


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    set_run_font(run, SIZE_PAGE)
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    run._r.extend([b, instr, e])


def add_toc(doc):
    title = doc.add_paragraph()
    set_paragraph_format(title, first_line=False)
    run = title.add_run("目录")
    set_run_font(run, SIZE_H2, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    run = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.extend([b, instr, sep])
    run2 = p.add_run()
    set_run_font(run2, SIZE_BODY)
    t = OxmlElement("w:t")
    t.text = "（请在 Word / WPS 中右键目录 → 更新域，以生成完整目录）"
    run2._r.append(t)
    run3 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run3._r.append(end)
    doc.add_page_break()


def enable_update_fields_on_open(doc):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is not None:
        settings.remove(existing)
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def peek_caption(lines, i, n):
    if i + 1 >= n:
        return None, i
    nxt = lines[i + 1].strip()
    m = re.match(r"^<p[^>]*>\s*<em>(.*?)</em>", nxt)
    if m:
        return m.group(1), i + 1
    m = re.match(r"^\*(图[：:].+|示意[：:].+)\*$", nxt)
    if m:
        return m.group(1), i + 1
    return None, i


def parse_md(doc, md_text: str, images: ImageIndex):
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    title_done = toc_inserted = False
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("<p") or stripped.startswith("<img"):
            block = [stripped]
            while i + 1 < n:
                joined = "\n".join(block)
                if "<img" in joined and joined.count("</p>") >= 1:
                    nxt = lines[i + 1].strip() if i + 1 < n else ""
                    if nxt.startswith("<p") and (
                        "<em>" in nxt or "示意" in nxt or "图" in nxt
                    ):
                        i += 1
                        block.append(lines[i].strip())
                    break
                i += 1
                if i >= n:
                    break
                block.append(lines[i].strip())
            joined = "\n".join(block)
            m = re.search(r'src="([^"]+)"', joined)
            cap_m = re.search(r"<em>(.*?)</em>", joined)
            if m:
                cap = cap_m.group(1) if cap_m else None
                if cap is None:
                    cap, i = peek_caption(lines, i, n)
                orig = m.group(1)
                display = images.numbered(orig)
                add_image(doc, display, images.resolve(orig), caption=cap)
            i += 1
            continue

        img_m = MD_IMG_RE.match(stripped)
        if img_m:
            cap, i = peek_caption(lines, i, n)
            orig = img_m.group(2)
            display = images.numbered(orig)
            add_image(doc, display, images.resolve(orig), caption=cap)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tl):
                    continue
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            add_table(doc, rows)
            continue

        if stripped.startswith("#"):
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level, text = len(m.group(1)), m.group(2).strip()
                if level == 1 and not title_done:
                    add_heading_styled(doc, text, 0)
                    title_done = True
                    if not toc_inserted:
                        add_toc(doc)
                        toc_inserted = True
                else:
                    add_heading_styled(doc, text, level - 1)
                i += 1
                continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            add_list_item(doc, m.group(2), ordered=True, index=int(m.group(1)))
            i += 1
            continue
        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:])
            i += 1
            continue

        if stripped.startswith("\\[") or stripped == "\\[" or stripped.startswith("$$"):
            buf = stripped
            closer = "$$" if stripped.startswith("$$") else "\\]"
            if closer not in stripped or stripped.count(closer) < (2 if closer == "$$" else 1):
                i += 1
                while i < n and closer not in lines[i]:
                    buf += " " + lines[i].strip()
                    i += 1
                if i < n:
                    buf += " " + lines[i].strip()
            latex = buf.replace("\\[", "").replace("\\]", "").replace("$$", "").strip()
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False)
            append_omml(p, latex, display=True)
            i += 1
            continue

        cap_line = stripped.strip("*").strip()
        if CAPTION_RE.match(cap_line) and (
            stripped.startswith("*") or cap_line.startswith(("图：", "图:", "示意：", "示意:"))
        ):
            add_caption(doc, cap_line)
            i += 1
            continue

        add_body_paragraph(doc, stripped)
        i += 1


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(NARROW_CM)
        section.left_margin = section.right_margin = Cm(NARROW_CM)
        add_page_number_footer(section)

    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = SIZE_BODY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    configure_heading_style(doc.styles["Heading 1"], SIZE_H1, Pt(18), Pt(8))
    configure_heading_style(doc.styles["Heading 2"], SIZE_H2, Pt(14), Pt(6))
    configure_heading_style(doc.styles["Heading 3"], SIZE_H3, Pt(10), Pt(6))
    configure_heading_style(doc.styles["Heading 4"], SIZE_H4, Pt(8), Pt(4))

    parse_md(doc, md_path.read_text(encoding="utf-8"), ImageIndex(md_path.parent))
    enable_update_fields_on_open(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main(argv=None):
    ap = argparse.ArgumentParser(description="Convert Markdown to formatted DOCX")
    ap.add_argument("input", type=Path, help="Input .md path")
    ap.add_argument("output", type=Path, help="Output .docx path")
    args = ap.parse_args(argv)
    if not args.input.exists():
        print(f"ERROR: input not found: {args.input}", file=sys.stderr)
        return 1
    try:
        import docx  # noqa: F401
    except ImportError:
        print("ERROR: pip install python-docx", file=sys.stderr)
        return 1
    convert(args.input, args.output)
    print(f"saved {args.output} ({args.output.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
